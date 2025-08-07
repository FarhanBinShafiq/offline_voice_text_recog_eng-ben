import tkinter as tk
from tkinter import filedialog
import sounddevice as sd
import numpy as np
import tempfile
import whisper
import time
import threading
import scipy.io.wavfile as wav
from docx import Document
import os

class WhisperStreamApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Whisper Multilingual Transcriber")
        self.root.configure(bg="#f4f4f4")

        self.model = whisper.load_model("medium")
        self.running = False
        self.sample_rate = 16000
        self.chunk_duration = 5

        self.language_map = {
            "English 🇺🇸": "en",
            "Hindi 🇮🇳": "hi",
            "Japanese 🇯🇵": "ja",
            "Spanish 🇪🇸": "es"
        }
        self.selected_language = tk.StringVar(value="English 🇺🇸")

        # Title label
        title_label = tk.Label(root, text="🎤 Whisper Multilingual Transcriber", font=("Helvetica", 18, "bold"), bg="#f4f4f4", fg="#333")
        title_label.pack(pady=(10, 0))

        self.text_area = tk.Text(root, wrap=tk.WORD, font=("Arial", 12), height=18, width=80, bg="white", bd=2, relief=tk.GROOVE)
        self.text_area.pack(padx=15, pady=10)

        button_frame = tk.Frame(root, bg="#f4f4f4")
        button_frame.pack(pady=10, fill="x")

        self.start_btn = tk.Button(button_frame, text="🎙 Start Listening", command=self.start_recording, bg="#4CAF50", fg="white", font=("Arial", 10, "bold"), padx=10, pady=5)
        self.start_btn.grid(row=0, column=0, padx=10)

        self.stop_btn = tk.Button(button_frame, text="🛑 Stop", command=self.stop_recording, state=tk.DISABLED, bg="#f44336", fg="white", font=("Arial", 10, "bold"), padx=10, pady=5)
        self.stop_btn.grid(row=0, column=1, padx=10)

        self.save_txt_btn = tk.Button(button_frame, text="💾 Save as .txt", command=self.save_as_txt, bg="#2196F3", fg="white", font=("Arial", 10), padx=10, pady=5)
        self.save_txt_btn.grid(row=0, column=2, padx=10)

        self.save_docx_btn = tk.Button(button_frame, text="📄 Save as .docx", command=self.save_as_docx, bg="#3f51b5", fg="white", font=("Arial", 10), padx=10, pady=5)
        self.save_docx_btn.grid(row=0, column=3, padx=10)

        lang_frame = tk.Frame(button_frame, bg="#f4f4f4")
        lang_frame.grid(row=1, column=0, columnspan=4, pady=(10, 0))
        tk.Label(lang_frame, text="Language:", bg="#f4f4f4").pack(side=tk.LEFT)
        lang_menu = tk.OptionMenu(lang_frame, self.selected_language, *self.language_map.keys())
        lang_menu.pack(side=tk.LEFT, padx=5)

        self.device_list = [dev['name'] for dev in sd.query_devices() if dev['max_input_channels'] > 0]
        self.device_index = tk.IntVar(value=0)

        mic_frame = tk.Frame(button_frame, bg="#f4f4f4")
        mic_frame.grid(row=2, column=0, columnspan=4, pady=5)
        tk.Label(mic_frame, text="Mic:", bg="#f4f4f4").pack(side=tk.LEFT)
        self.device_menu = tk.OptionMenu(mic_frame, self.device_index, *range(len(self.device_list)))
        self.device_menu.pack(side=tk.LEFT, padx=5)

        self.device_label = tk.Label(mic_frame, text=self.device_list[0][:40], bg="#f4f4f4")
        self.device_label.pack(side=tk.LEFT, padx=10)

        self.device_index.trace("w", self.update_device_label)

    def update_device_label(self, *args):
        i = self.device_index.get()
        if 0 <= i < len(self.device_list):
            self.device_label.config(text=self.device_list[i][:60])

    def start_recording(self):
        if not self.running:
            self.running = True
            self.start_btn.config(state=tk.DISABLED)
            self.stop_btn.config(state=tk.NORMAL)
            threading.Thread(target=self.record_loop, daemon=True).start()
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Started listening...\n")

    def stop_recording(self):
        self.running = False
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Stopped.\n")

    def record_loop(self):
        while self.running:
            try:
                sd.default.device = self.device_index.get()
                audio = sd.rec(int(self.chunk_duration * self.sample_rate),
                               samplerate=self.sample_rate,
                               channels=1,
                               dtype='int16')
                sd.wait()

                if audio is None or audio.size == 0:
                    self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] [SKIPPED] No audio captured.\n")
                    continue

                threading.Thread(target=self.transcribe_chunk, args=(audio.copy(),), daemon=True).start()
            except Exception as e:
                self.text_area.insert(tk.END, f"[ERROR] Failed to record: {e}\n")

    def transcribe_chunk(self, audio):
        audio = np.squeeze(audio.astype(np.float32))

        temp_dir = os.path.join(os.getcwd(), "temp_audio")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, f"chunk_{int(time.time() * 1000)}.wav")

        try:
            wav.write(temp_path, self.sample_rate, audio.astype(np.int16))
            lang_code = self.language_map[self.selected_language.get()]
            result = self.model.transcribe(temp_path, language=lang_code)
            text = result.get("text", "").strip()

            if text:
                timestamp = time.strftime('%H:%M:%S')
                self.text_area.insert(tk.END, f"[{timestamp}] {text}\n")
                self.text_area.see(tk.END)
        except Exception as e:
            self.text_area.insert(tk.END, f"[ERROR] {str(e)}\n")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def save_as_txt(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if file_path:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(self.text_area.get("1.0", tk.END))
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Saved as .txt\n")

    def save_as_docx(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if file_path:
            doc = Document()
            doc.add_heading("Whisper Transcript", level=1)
            for line in self.text_area.get("1.0", tk.END).split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Saved as .docx\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = WhisperStreamApp(root)
    root.mainloop()
