import tkinter as tk
from tkinter import filedialog
import sounddevice as sd
import numpy as np
import tempfile
import whisper
import time
import threading
import scipy.io.wavfile as wav
from docx import Document
import os


class WhisperStreamApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Whisper Stream Transcriber (English Only)")

        self.model = whisper.load_model("base.en")  # Optimized for English
        self.running = False
        self.sample_rate = 16000
        self.chunk_duration = 5  # seconds per chunk

        # Transcript display
        self.text_area = tk.Text(root, wrap=tk.WORD, font=("Arial", 12), height=20, width=70)
        self.text_area.pack(padx=10, pady=10)

        # Controls
        button_frame = tk.Frame(root)
        button_frame.pack(pady=10)

        self.start_btn = tk.Button(button_frame, text="🎙 Start Listening", command=self.start_recording)
        self.start_btn.grid(row=0, column=0, padx=10)

        self.stop_btn = tk.Button(button_frame, text="🛑 Stop", command=self.stop_recording, state=tk.DISABLED)
        self.stop_btn.grid(row=0, column=1, padx=10)

        self.save_txt_btn = tk.Button(button_frame, text="💾 Save as .txt", command=self.save_as_txt)
        self.save_txt_btn.grid(row=0, column=2, padx=10)

        self.save_docx_btn = tk.Button(button_frame, text="📄 Save as .docx", command=self.save_as_docx)
        self.save_docx_btn.grid(row=0, column=3, padx=10)

        # Mic selector
        self.device_list = [dev['name'] for dev in sd.query_devices() if dev['max_input_channels'] > 0]
        self.device_index = tk.IntVar(value=0)

        tk.Label(button_frame, text="Mic Device").grid(row=1, column=0, columnspan=2)
        self.device_menu = tk.OptionMenu(button_frame, self.device_index, *range(len(self.device_list)))
        self.device_menu.grid(row=1, column=2, columnspan=2)

        self.device_label = tk.Label(button_frame, text=self.device_list[0][:40])
        self.device_label.grid(row=2, column=0, columnspan=5, pady=5)

        self.device_index.trace("w", self.update_device_label)

    def update_device_label(self, *args):
        i = self.device_index.get()
        if 0 <= i < len(self.device_list):
            self.device_label.config(text=self.device_list[i][:60])

    def start_recording(self):
        if not self.running:
            self.running = True
            self.start_btn.config(state=tk.DISABLED)
            self.stop_btn.config(state=tk.NORMAL)
            threading.Thread(target=self.record_loop, daemon=True).start()
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Started listening...\n")

    def stop_recording(self):
        self.running = False
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Stopped.\n")

    def record_loop(self):
        while self.running:
            try:
                sd.default.device = self.device_index.get()
                audio = sd.rec(int(self.chunk_duration * self.sample_rate),
                               samplerate=self.sample_rate,
                               channels=1,
                               dtype='int16')
                sd.wait()

                if audio is None or audio.size == 0:
                    self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] [SKIPPED] No audio captured.\n")
                    continue

                threading.Thread(target=self.transcribe_chunk, args=(audio.copy(),), daemon=True).start()
            except Exception as e:
                self.text_area.insert(tk.END, f"[ERROR] Failed to record: {e}\n")

    def transcribe_chunk(self, audio):
        if audio is None or audio.size == 0:
            return

        temp_dir = os.path.join(os.getcwd(), "temp_audio")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, f"chunk_{int(time.time() * 1000)}.wav")

        try:
            wav.write(temp_path, self.sample_rate, audio)
            result = self.model.transcribe(temp_path)  # Language is now always English
            text = result.get("text", "").strip()

            if text:
                timestamp = time.strftime('%H:%M:%S')
                self.text_area.insert(tk.END, f"[{timestamp}] {text}\n")
                self.text_area.see(tk.END)

        except Exception as e:
            self.text_area.insert(tk.END, f"[ERROR] {str(e)}\n")

        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def save_as_txt(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if file_path:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(self.text_area.get("1.0", tk.END))
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Saved as .txt\n")

    def save_as_docx(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if file_path:
            doc = Document()
            doc.add_heading("Whisper Transcript", level=1)
            for line in self.text_area.get("1.0", tk.END).split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
            self.text_area.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] Saved as .docx\n")


# Run it
if __name__ == "__main__":
    root = tk.Tk()
    app = WhisperStreamApp(root)
    root.mainloop()
