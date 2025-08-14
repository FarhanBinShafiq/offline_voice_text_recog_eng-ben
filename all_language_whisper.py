import tkinter as tk
from tkinter import ttk, filedialog
import numpy as np
import sounddevice as sd
import threading, queue, time, os, json, datetime, collections

from docx import Document
from faster_whisper import WhisperModel

# ===== Optional extras =====
try:
    import soundcard as sc
    HAS_SOUNDCARD = True
except Exception:
    HAS_SOUNDCARD = False

try:
    import noisereduce as nr
    HAS_NOISEREDUCE = True
except Exception:
    HAS_NOISEREDUCE = False

try:
    import webrtcvad
    HAS_WEBRTCVAD = True
except Exception:
    HAS_WEBRTCVAD = False


# ---------- Small utils ----------
def ts(): return time.strftime("%H:%M:%S")
def now_tag(): return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
def ensure_dir(path): os.makedirs(path, exist_ok=True); return path


def get_default_input_device_index() -> int:
    """Default input device index, fallback to first input-capable device."""
    try:
        d = sd.default.device
        if isinstance(d, (list, tuple)) and d[0] is not None:
            return int(d[0])
        if isinstance(d, int):
            return d
    except Exception:
        pass
    for i, dev in enumerate(sd.query_devices()):
        if dev.get("max_input_channels", 0) > 0:
            return i
    raise RuntimeError("No input microphone found")


def best_samplerate_for_device(device_index: int) -> int:
    """Prefer 16k for ASR; fallback to 8k; else host default."""
    for sr in (16000, 8000):
        try:
            sd.check_input_settings(device=device_index, samplerate=sr, channels=1)
            return sr
        except Exception:
            continue
    info = sd.query_devices(device_index)
    host = sd.query_hostapis()[info["hostapi"]]
    default_sr = int(host.get("default_samplerate", 16000) or 16000)
    sd.check_input_settings(device=device_index, samplerate=default_sr, channels=1)
    return default_sr


def linear_resample_1d(x: np.ndarray, sr_in: int, sr_out: int) -> np.ndarray:
    """Lightweight linear resampler (mono float32)."""
    if sr_in == sr_out or x.size == 0:
        return x.astype(np.float32, copy=False)
    ratio = float(sr_out) / float(sr_in)
    new_len = max(1, int(round(x.size * ratio)))
    xi = np.linspace(0.0, 1.0, x.size, dtype=np.float32)
    xo = np.linspace(0.0, 1.0, new_len, dtype=np.float32)
    return np.interp(xo, xi, x.astype(np.float32)).astype(np.float32)


def float_to_int16(x: np.ndarray) -> np.ndarray:
    x = np.clip(x, -1.0, 1.0)
    return (x * 32767.0).astype(np.int16)


def int16_to_float(x: np.ndarray) -> np.ndarray:
    return (x.astype(np.float32) / 32768.0)


# ---------- Ring buffer ----------
class RingBuffer:
    def __init__(self, max_chunks=256):
        self.q = queue.Queue(maxsize=max_chunks)

    def push(self, arr: np.ndarray):
        try:
            self.q.put_nowait(arr)
        except queue.Full:
            # drop oldest
            try:
                _ = self.q.get_nowait()
                self.q.put_nowait(arr)
            except Exception:
                pass

    def pop_all(self):
        items = []
        try:
            while True:
                items.append(self.q.get_nowait())
        except queue.Empty:
            pass
        if not items:
            return None
        return np.concatenate(items, axis=0)


# ---------- Preprocessing ----------
class Preprocessor:
    """Noise reduction + VAD (optional). Works on 16k mono."""
    def __init__(self, use_nr=False, use_vad=False, vad_aggressiveness=2):
        self.use_nr = use_nr and HAS_NOISEREDUCE
        self.use_vad = use_vad and HAS_WEBRTCVAD
        self.vad = webrtcvad.Vad(vad_aggressiveness) if self.use_vad else None

    def apply(self, audio_16k_f32: np.ndarray) -> np.ndarray:
        y = audio_16k_f32
        if self.use_nr:
            try:
                # Fast stationary noise reduction (keeps latency small)
                y = nr.reduce_noise(y=y, sr=16000, stationary=True, prop_decrease=0.9, use_tensorflow=False)
            except Exception:
                pass

        if self.use_vad:
            try:
                # WebRTC VAD expects 16k mono int16 frames of 10/20/30 ms
                frame_ms = 30
                frame_len = int(16000 * frame_ms / 1000)
                i16 = float_to_int16(y)
                out = []
                for i in range(0, len(i16), frame_len):
                    frame = i16[i:i+frame_len]
                    if len(frame) < frame_len:
                        break
                    if self.vad.is_speech(frame.tobytes(), 16000):
                        out.append(frame)
                if out:
                    y = int16_to_float(np.concatenate(out))
                else:
                    # if VAD drops everything, keep tiny noise to avoid empty input
                    y = np.zeros(1, dtype=np.float32)
            except Exception:
                pass

        return y.astype(np.float32, copy=False)


# ---------- ASR worker ----------
class ASRWorker(threading.Thread):
    """
    Consumes int16 audio at any SR, resamples to 16k, preprocess (NR/VAD),
    transcribes with faster-whisper. Smart auto-language stabilization.
    """
    def __init__(self, rb, on_text, on_info, stop_event,
                 model_name="large-v3", compute_type="float16",
                 language=None, task="transcribe",
                 input_sr=16000, chunk_secs=4, overlap_secs=0.5,
                 use_nr=False, use_vad=False, smart_lang=True):
        super().__init__(daemon=True)
        self.rb = rb
        self.on_text = on_text
        self.on_info = on_info
        self.stop_event = stop_event

        self.model_name = model_name
        self.compute_type = compute_type
        self.user_language = language  # None = auto
        self.smart_lang = smart_lang   # dynamic stabilization if auto
        self.task = task

        self.input_sr = input_sr
        self.target_sr = 16000
        self.stride_in = int(chunk_secs * input_sr)
        self.overlap_in = int(overlap_secs * input_sr)

        self.tail = np.zeros((0, 1), dtype=np.int16)
        self.model = None

        self.pp = Preprocessor(use_nr=use_nr, use_vad=use_vad, vad_aggressiveness=2)

        # Language state (only used when smart auto)
        self.lang_window = collections.deque(maxlen=8)  # (lang, prob)
        self.locked_language = None

    def _decide_language(self):
        """Majority + confidence heuristic to stabilize language in auto mode."""
        if not self.lang_window:
            return None
        counts = {}
        probs = {}
        for lang, prob in self.lang_window:
            counts[lang] = counts.get(lang, 0) + 1
            probs[lang] = max(prob, probs.get(lang, 0))
        # pick lang with highest count, break ties with max prob
        best = None
        for lang in counts:
            cand = (counts[lang], probs[lang], lang)
            if best is None or cand > best:
                best = cand
        count, prob, lang = best
        # lock if confident enough
        if count >= 3 and prob >= 0.80:
            return lang
        return None

    def run(self):
        try:
            # tune workers: more workers can help on GPU or fast CPU
            self.model = WhisperModel(
                self.model_name, device="auto",
                compute_type=self.compute_type,
                cpu_threads=0,            # let lib decide optimal
                num_workers=2
            )
            self.on_info(f"[{ts()}] ASR loaded: {self.model_name} ({self.compute_type})")
            self.on_info(f"[{ts()}] Input SR: {self.input_sr} Hz → ASR SR: {self.target_sr} Hz")
        except Exception as e:
            self.on_info(f"[{ts()}] [ERROR] Model load failed: {e}")
            return

        while not self.stop_event.is_set():
            time.sleep(0.01)
            buf = self.rb.pop_all()
            if buf is None or buf.size == 0:
                continue

            if self.tail.size:
                buf = np.concatenate([self.tail, buf], axis=0)

            if buf.shape[0] < self.stride_in:
                self.tail = buf
                continue

            chunk = buf[:self.stride_in]
            keep_from = max(0, self.stride_in - self.overlap_in)
            self.tail = buf[keep_from:]

            # int16 -> float32 [-1,1] @ input_sr
            audio_f32 = int16_to_float(chunk.squeeze(axis=-1))
            # resample to 16k for ASR
            audio_16k = linear_resample_1d(audio_f32, self.input_sr, self.target_sr)
            # preprocess
            audio_16k = self.pp.apply(audio_16k)

            # Determine language
            lang_for_this = self.user_language
            if self.user_language is None and self.smart_lang:
                # use unlocked language first; let info decide to lock later
                lang_for_this = self.locked_language  # may be None

            try:
                segments, info = self.model.transcribe(
                    audio_16k,
                    language=lang_for_this,             # None = auto
                    task=self.task,                     # 'transcribe' or 'translate'
                    vad_filter=False,                   # using our own VAD
                    beam_size=1, best_of=1,
                    no_speech_threshold=0.2,
                    condition_on_previous_text=False,   # lowers latency, avoids bias
                    word_timestamps=False,
                    temperature=0.0,
                )

                # Update language window if in auto/smart mode
                if self.user_language is None and self.smart_lang and info is not None:
                    det_lang = getattr(info, "language", None) or (info.get("language") if isinstance(info, dict) else None)
                    det_prob = float(getattr(info, "language_probability", 0.0) or (info.get("language_probability", 0.0) if isinstance(info, dict) else 0.0))
                    if det_lang:
                        self.lang_window.append((det_lang, det_prob))
                        new_lock = self._decide_language()
                        if new_lock and new_lock != self.locked_language:
                            self.locked_language = new_lock
                            self.on_info(f"[{ts()}] 🔒 Language locked: {self.locked_language} (auto)")

                out = []
                srt_segments = []
                for seg in segments:
                    text = (seg.text or "").strip()
                    if text:
                        out.append(text)
                        srt_segments.append({
                            "start": float(seg.start or 0.0),
                            "end": float(seg.end or 0.0),
                            "text": text
                        })
                if out:
                    self.on_text(" ".join(out), srt_segments)
            except Exception as e:
                self.on_info(f"[{ts()}] [ERROR] ASR failed: {e}")


# ---------- Languages ----------
LANG_OPTIONS = [
    "Auto (smart detect)",
    "Afrikaans (af)", "Albanian (sq)", "Amharic (am)", "Arabic (ar)", "Armenian (hy)",
    "Assamese (as)", "Azerbaijani (az)", "Bashkir (ba)", "Basque (eu)", "Belarusian (be)",
    "Bengali (bn)", "Bosnian (bs)", "Breton (br)", "Bulgarian (bg)", "Burmese (my)",
    "Catalan (ca)", "Chinese (zh)", "Cantonese (yue)", "Croatian (hr)", "Czech (cs)", "Danish (da)",
    "Dutch (nl)", "English (en)", "Estonian (et)", "Faroese (fo)", "Finnish (fi)",
    "French (fr)", "Galician (gl)", "Georgian (ka)", "German (de)", "Greek (el)",
    "Gujarati (gu)", "Haitian Creole (ht)", "Hausa (ha)", "Hawaiian (haw)", "Hebrew (he)",
    "Hindi (hi)", "Hungarian (hu)", "Icelandic (is)", "Indonesian (id)", "Italian (it)",
    "Japanese (ja)", "Javanese (jv)", "Kannada (kn)", "Kazakh (kk)", "Khmer (km)",
    "Korean (ko)", "Lao (lo)", "Latin (la)", "Latvian (lv)", "Lingala (ln)",
    "Lithuanian (lt)", "Luxembourgish (lb)", "Macedonian (mk)", "Malagasy (mg)", "Malay (ms)",
    "Malayalam (ml)", "Maltese (mt)", "Maori (mi)", "Marathi (mr)", "Mongolian (mn)",
    "Nepali (ne)", "Norwegian (no)", "Occitan (oc)", "Oriya (or)", "Oromo (om)",
    "Pashto (ps)", "Persian (fa)", "Polish (pl)", "Portuguese (pt)", "Punjabi (pa)",
    "Romanian (ro)", "Russian (ru)", "Sanskrit (sa)", "Scottish Gaelic (gd)", "Serbian (sr)",
    "Shona (sn)", "Sindhi (sd)", "Sinhala (si)", "Slovak (sk)", "Slovenian (sl)",
    "Somali (so)", "Spanish (es)", "Sundanese (su)", "Swahili (sw)", "Swedish (sv)",
    "Tagalog (tl)", "Tajik (tg)", "Tamil (ta)", "Tatar (tt)", "Telugu (te)",
    "Thai (th)", "Turkish (tr)", "Turkmen (tk)", "Ukrainian (uk)", "Urdu (ur)",
    "Uzbek (uz)", "Vietnamese (vi)", "Welsh (cy)", "Yiddish (yi)", "Yoruba (yo)", "Zulu (zu)",
    "Hmong (hmn)", "Kurdish (ku)", "Kurdish Sorani (ckb)", "Uighur (ug)",
]


# ---------- GUI ----------
class MeetingTranscriberApp:
    DEFAULT_SR = 16000  # target SR for ASR (resampled)

    def __init__(self, root):
        self.root = root
        self.root.title("Meeting Transcriber (Offline)")
        self.root.geometry("1060x740")
        self.root.configure(bg="#0b1220")

        # state
        self.running = False
        self.stop_event = threading.Event()
        self.rb = RingBuffer()

        self.lines = []     # [{"time": "...", "text": "..."}]
        self.segments = []  # [{"start": s, "end": e, "text": "..."}]
        self.last_autosave = time.time()
        self.meeting_dir = ensure_dir(os.path.join(os.getcwd(), f"meeting_{now_tag()}"))

        # UI state
        self.model_name = tk.StringVar(value="large-v3")
        self.compute_type = tk.StringVar(value="float16")  # auto-downgrades on CPU
        self.lang = tk.StringVar(value="Auto (smart detect)")
        self.translate = tk.BooleanVar(value=False)
        self.chunk_secs = tk.IntVar(value=4)
        self.overlap = tk.DoubleVar(value=0.5)
        self.meeting_mode = tk.BooleanVar(value=False)
        self.use_nr = tk.BooleanVar(value=True and HAS_NOISEREDUCE)
        self.use_vad = tk.BooleanVar(value=True and HAS_WEBRTCVAD)
        self.smart_lang = tk.BooleanVar(value=True)

        # capture settings (resolved at start)
        self.mic_device_index = None
        self.input_sr = self.DEFAULT_SR

        self._build_ui()

    # ----- UI -----
    def _build_ui(self):
        s = ttk.Style()
        try: s.theme_use("clam")
        except Exception: pass
        s.configure("TFrame", background="#0b1220")
        s.configure("Card.TFrame", background="#0f1626", relief="groove", borderwidth=1)
        s.configure("TLabel", foreground="#e5e7eb", background="#0b1220")
        s.configure("Title.TLabel", font=("Segoe UI", 18, "bold"), foreground="#e5e7eb", background="#0b1220")
        s.configure("Sub.TLabel", font=("Segoe UI", 10), foreground="#9ca3af", background="#0b1220")
        s.configure("TButton", font=("Segoe UI", 10))
        s.configure("Accent.TButton", font=("Segoe UI", 10, "bold"))
        s.map("Accent.TButton", background=[("!disabled", "#22c55e"), ("disabled", "#14532d")], foreground=[("!disabled", "white")])
        s.configure("Danger.TButton", font=("Segoe UI", 10, "bold"))
        s.map("Danger.TButton", background=[("!disabled", "#ef4444"), ("disabled", "#7f1d1d")], foreground=[("!disabled", "white")])
        s.configure("TCheckbutton", background="#0b1220", foreground="#e5e7eb")

        header = ttk.Frame(self.root, style="TFrame")
        header.pack(fill="x", padx=16, pady=(16, 8))
        ttk.Label(header, text="🎤 Meeting Transcriber (Offline)", style="Title.TLabel").pack(side="left")
        ttk.Label(header, text="Faster-Whisper • Smart Auto-Language • VAD/Noise Reduction", style="Sub.TLabel").pack(side="right")

        card = ttk.Frame(self.root, style="Card.TFrame")
        card.pack(fill="x", padx=16, pady=8)

        row1 = ttk.Frame(card, style="TFrame")
        row1.pack(fill="x", padx=12, pady=10)
        self.btn_start = ttk.Button(row1, text="Start", style="Accent.TButton", command=self.start)
        self.btn_stop  = ttk.Button(row1, text="Stop", style="Danger.TButton", state="disabled", command=self.stop)
        self.btn_clear = ttk.Button(row1, text="Clear", command=self.clear)
        self.btn_save_txt  = ttk.Button(row1, text="Save .txt", command=self.save_txt)
        self.btn_save_docx = ttk.Button(row1, text="Save .docx", command=self.save_docx)
        self.btn_save_json = ttk.Button(row1, text="Save .json", command=self.save_json)
        self.btn_save_srt  = ttk.Button(row1, text="Save .srt", command=self.save_srt)
        for w in (self.btn_start, self.btn_stop, self.btn_clear, self.btn_save_txt, self.btn_save_docx, self.btn_save_json, self.btn_save_srt):
            w.pack(side="left", padx=6)

        row2 = ttk.Frame(card, style="TFrame"); row2.pack(fill="x", padx=12, pady=6)
        ttk.Label(row2, text="Model").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.model_name, values=["base","small","medium","large-v3"], width=12, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Label(row2, text="Compute").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.compute_type, values=["int8","float16","float32"], width=10, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Label(row2, text="Language").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.lang, values=LANG_OPTIONS, width=24, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row2, text="Translate to English", variable=self.translate).pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row2, text=f"Meeting Mode (capture speakers){'' if HAS_SOUNDCARD else ' [install soundcard]'}",
                        variable=self.meeting_mode, state=("!disabled" if HAS_SOUNDCARD else "disabled")).pack(side="left")

        row3 = ttk.Frame(card, style="TFrame"); row3.pack(fill="x", padx=12, pady=(0, 8))
        ttk.Label(row3, text="Chunk (s)").pack(side="left", padx=(0, 6))
        self.scale_chunk = ttk.Scale(row3, from_=2, to=8, orient="horizontal")
        self.scale_chunk.pack(side="left", padx=(0, 16))
        self.lbl_chunk = ttk.Label(row3, text=f"{self.chunk_secs.get()}s")
        self.lbl_chunk.pack(side="left", padx=(0, 16))
        self.scale_chunk.set(self.chunk_secs.get())
        self.scale_chunk.configure(command=self._on_chunk_change)

        ttk.Label(row3, text="Overlap").pack(side="left", padx=(0, 6))
        self.scale_olap = ttk.Scale(row3, from_=0.3, to=0.8, orient="horizontal")
        self.scale_olap.pack(side="left", padx=(0, 16))
        self.lbl_olap = ttk.Label(row3, text=f"{self.overlap.get():.1f}")
        self.lbl_olap.pack(side="left", padx=(0, 16))
        self.scale_olap.set(self.overlap.get())
        self.scale_olap.configure(command=self._on_overlap_change)

        row4 = ttk.Frame(card, style="TFrame"); row4.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Checkbutton(row4, text=f"Noise reduction ({'on' if self.use_nr.get() else 'off'})", variable=self.use_nr).pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row4, text=f"WebRTC VAD ({'on' if self.use_vad.get() else 'off'})", variable=self.use_vad).pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row4, text="Smart auto-language", variable=self.smart_lang).pack(side="left", padx=(0, 16))

        transcript = ttk.Frame(self.root, style="Card.TFrame")
        transcript.pack(fill="both", expand=True, padx=16, pady=(8, 16))
        ttk.Label(transcript, text="Transcript", style="Sub.TLabel").pack(anchor="w", padx=12, pady=(8, 0))
        self.text = tk.Text(transcript, wrap="word", font=("Segoe UI", 11),
                            bg="#09101d", fg="#e5e7eb", insertbackground="#e5e7eb", relief="flat")
        self.text.pack(fill="both", expand=True, padx=12, pady=12)

        footer = ttk.Frame(self.root, style="TFrame"); footer.pack(fill="x", padx=16, pady=(0, 12))
        self.status = ttk.Label(footer, text="Ready.", style="Sub.TLabel"); self.status.pack(side="left")

        # Availability hints
        if not HAS_NOISEREDUCE:
            self._append_line("[info] noisereduce not installed — Noise Reduction toggle will be ignored.")
        if not HAS_WEBRTCVAD:
            self._append_line("[info] webrtcvad not installed — VAD toggle will be ignored.")

    def _on_chunk_change(self, *_):
        v = int(float(self.scale_chunk.get())); self.chunk_secs.set(v)
        if hasattr(self, "lbl_chunk"): self.lbl_chunk.config(text=f"{v}s")

    def _on_overlap_change(self, *_):
        v = float(self.scale_olap.get()); self.overlap.set(round(v, 1))
        if hasattr(self, "lbl_olap"): self.lbl_olap.config(text=f"{self.overlap.get():.1f}")

    def set_status(self, msg): self.status.config(text=msg)

    # ----- Start/Stop -----
    def start(self):
        if self.running: return
        self.running = True
        self.stop_event.clear()
        self.btn_start.config(state="disabled"); self.btn_stop.config(state="normal")
        self.lines.clear(); self.segments.clear()
        self.set_status("Starting…")

        # resolve mic + SR (for mic mode)
        try:
            self.mic_device_index = get_default_input_device_index()
            self.input_sr = best_samplerate_for_device(self.mic_device_index)
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Mic detect: {e}")
            self.stop(); return

        # language / task
        selected = self.lang.get()
        lang = None
        smart = self.smart_lang.get()
        if selected and not selected.lower().startswith("auto"):
            lang = selected.split("(")[-1].split(")")[0].strip()
            smart = False  # explicit language → no smart auto
        task = "translate" if self.translate.get() else "transcribe"

        # ASR worker
        self.asr_thread = ASRWorker(
            rb=self.rb, on_text=self._on_text, on_info=self._append_line, stop_event=self.stop_event,
            model_name=self.model_name.get(), compute_type=self.compute_type.get(),
            language=lang, task=task,
            input_sr=(self.input_sr if not self.meeting_mode.get() else 48000),
            chunk_secs=self.chunk_secs.get(), overlap_secs=self.overlap.get(),
            use_nr=self.use_nr.get(), use_vad=self.use_vad.get(), smart_lang=smart
        )
        self.asr_thread.start()

        # recorder thread
        rec_target = self._record_loop_meeting if (self.meeting_mode.get() and HAS_SOUNDCARD) else self._record_loop_mic
        self.rec_thread = threading.Thread(target=rec_target, daemon=True)
        self.rec_thread.start()
        self._append_line(f"[{ts()}] Meeting started. Saving in: {self.meeting_dir}")

    def stop(self):
        if not self.running: return
        self.running = False
        self.stop_event.set()
        self.btn_start.config(state="normal"); self.btn_stop.config(state="disabled")
        self.set_status("Stopping…"); time.sleep(0.3); self.set_status("Stopped.")

    def clear(self):
        self.text.delete("1.0", tk.END)
        self.lines.clear(); self.segments.clear()

    # ----- Recording -----
    def _record_loop_mic(self):
        """Capture laptop/default mic at the device's supported SR (8k/16k/etc)."""
        frames = 512
        try:
            stream = sd.InputStream(samplerate=self.input_sr, channels=1, dtype="int16",
                                    blocksize=frames, device=self.mic_device_index)
            stream.start()
            self._append_line(f"[{ts()}] Mic capture: device #{self.mic_device_index} @ {self.input_sr} Hz")
            self.set_status("Listening (mic)…")

            accum = []
            target = int(self.chunk_secs.get() * self.input_sr)
            while not self.stop_event.is_set():
                data, _ = stream.read(frames)
                if data is not None and data.size > 0:
                    accum.append(data)
                total = sum(x.shape[0] for x in accum)
                if total >= max(256, target // 6):  # push more frequently for latency
                    self.rb.push(np.concatenate(accum, axis=0)); accum = []
                self._autosave_if_needed()
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Mic: {e}")
        finally:
            try: stream.stop(); stream.close()
            except Exception: pass

    def _record_loop_meeting(self):
        """Capture system audio (speakers) + mic mix using soundcard at 48 kHz, then queue int16."""
        if not HAS_SOUNDCARD:
            self._append_line(f"[{ts()}] [WARN] soundcard not installed; fallback to mic.")
            self._record_loop_mic(); return

        sr = 48000
        block = 512
        try:
            speaker = sc.default_speaker(); mic = sc.default_microphone()
            self._append_line(f"[{ts()}] Meeting mode: loopback @ {sr} Hz")
            self.set_status("Listening (meeting mode)…")
            with mic.recorder(samplerate=sr, channels=1, blocksize=block) as mic_rec, \
                 speaker.recorder(samplerate=sr, channels=1, blocksize=block) as spk_rec:
                accum = []
                target = int(self.chunk_secs.get() * sr)
                while not self.stop_event.is_set():
                    mic_data = mic_rec.record(numframes=block)      # float32 [-1,1]
                    spk_data = spk_rec.record(numframes=block)      # float32 [-1,1]
                    mixed = np.clip(mic_data + spk_data, -1.0, 1.0) # mix
                    i16 = float_to_int16(mixed.squeeze(-1))[:, None]  # → int16, keep (N,1)
                    accum.append(i16)
                    total = sum(x.shape[0] for x in accum)
                    if total >= max(768, target // 6):
                        self.rb.push(np.concatenate(accum, axis=0)); accum = []
                    self._autosave_if_needed()
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Meeting mode: {e}")
            self._record_loop_mic()

    # ----- Callbacks from ASR -----
    def _on_text(self, text_line, srt_segments):
        self.text.insert(tk.END, f"[{ts()}] {text_line}\n")
        self.text.see(tk.END)
        self.lines.append({"time": ts(), "text": text_line})
        self.segments.extend(srt_segments)
        self._autosave_if_needed()

    def _append_line(self, text):
        self.text.insert(tk.END, text + "\n")
        self.text.see(tk.END)

    # ----- Autosave -----
    def _autosave_if_needed(self):
        if time.time() - self.last_autosave > 60:
            self._save_txt(os.path.join(self.meeting_dir, "autosave.txt"))
            self._save_json(os.path.join(self.meeting_dir, "autosave.json"))
            self._save_srt(os.path.join(self.meeting_dir, "autosave.srt"))
            self.last_autosave = time.time()

    # ----- Save actions -----
    def save_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not path: return
        self._save_txt(path)

    def save_docx(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not path: return
        doc = Document(); doc.add_heading("Meeting Transcript", level=1)
        for line in self.lines:
            doc.add_paragraph(f"[{line['time']}] {line['text']}")
        doc.save(path); self.set_status(f"Saved {os.path.basename(path)}")

    def save_json(self):
        path = filedialog.asksaveasfilename(defaultextension=".json")
        if not path: return
        self._save_json(path)

    def save_srt(self):
        path = filedialog.asksaveasfilename(defaultextension=".srt")
        if not path: return
        self._save_srt(path)

    # ----- Save helpers -----
    def _save_txt(self, path):
        with open(path, "w", encoding="utf-8") as f:
            for line in self.lines:
                f.write(f"[{line['time']}] {line['text']}\n")
        self.set_status(f"Saved {os.path.basename(path)}")

    def _save_json(self, path):
        payload = {
            "created": datetime.datetime.now().isoformat(),
            "model": self.model_name.get(),
            "compute": self.compute_type.get(),
            "language": self.lang.get(),
            "translate_to_english": self.translate.get(),
            "lines": self.lines,
            "segments": self.segments,
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        self.set_status(f"Saved {os.path.basename(path)}")

    def _save_srt(self, path):
        def fmt_time(sec):
            h = int(sec // 3600); m = int((sec % 3600) // 60); s = int(sec % 60)
            ms = int((sec - int(sec)) * 1000)
            return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

        out, idx = [], 1
        for seg in self.segments:
            start = max(0.0, float(seg.get("start", 0.0)))
            end   = max(start, float(seg.get("end", start + 0.1)))
            text  = (seg.get("text") or "").strip()
            if not text: continue
            out.append(f"{idx}\n{fmt_time(start)} --> {fmt_time(end)}\n{text}\n")
            idx += 1
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(out))
        self.set_status(f"Saved {os.path.basename(path)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = MeetingTranscriberApp(root)
    root.mainloop()
