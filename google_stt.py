import tkinter as tk
from tkinter import ttk, filedialog
import numpy as np
import sounddevice as sd
import threading, queue, time, os, json, datetime
from docx import Document

# Google Cloud Speech-to-Text
try:
    from google.cloud import speech as gspeech
except Exception as e:
    gspeech = None
    print("google-cloud-speech is not installed. Run: pip install google-cloud-speech")

# ---------------- Small utils ----------------
def ts(): return time.strftime("%H:%M:%S")
def now_tag(): return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
def ensure_dir(path): os.makedirs(path, exist_ok=True); return path

TARGET_SR = 16000  # Google STT streaming best practice

def pick_input_device_and_sr():
    """Pick a mic device and a workable sample rate (try 16k first)."""
    dev_index = None
    try:
        d = sd.default.device
        if isinstance(d, (list, tuple)) and d[0] is not None:
            dev_index = int(d[0])
    except Exception:
        pass

    if dev_index is None:
        for i, dev in enumerate(sd.query_devices()):
            if dev.get("max_input_channels", 0) > 0:
                dev_index = i
                break
    if dev_index is None:
        raise RuntimeError("No input microphone found")

    # Try preferred SRs (16k first for STT)
    for sr in (16000, 48000, 44100, 8000):
        try:
            sd.check_input_settings(device=dev_index, samplerate=sr, channels=1)
            return dev_index, sr
        except Exception:
            continue

    info = sd.query_devices(dev_index)
    default_sr = int(info.get("default_samplerate", 16000) or 16000)
    sd.check_input_settings(device=dev_index, samplerate=default_sr, channels=1)
    return dev_index, default_sr

def linear_resample_1d(x: np.ndarray, sr_in: int, sr_out: int) -> np.ndarray:
    """Very simple linear resampler; x is mono float32."""
    if sr_in == sr_out or x.size == 0:
        return x.astype(np.float32, copy=False)
    ratio = float(sr_out) / float(sr_in)
    new_len = max(1, int(round(x.size * ratio)))
    xi = np.linspace(0.0, 1.0, x.size, dtype=np.float32)
    xo = np.linspace(0.0, 1.0, new_len, dtype=np.float32)
    return np.interp(xo, xi, x.astype(np.float32)).astype(np.float32)

# ---------------- App ----------------
class GoogleSTTApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Simple STT (Google Cloud)")
        self.root.geometry("880x620")

        # state
        self.running = False
        self.stop_event = threading.Event()
        self.audio_q = queue.Queue(maxsize=100)
        self.lines = []   # [{"time": "...", "text": "..."}]
        self.meeting_dir = ensure_dir(os.path.join(os.getcwd(), f"stt_{now_tag()}"))

        # mic
        self.mic_device_index = None
        self.input_sr = TARGET_SR

        # UI state
        self.primary_lang = tk.StringVar(value="bn-BD")         # default: Bengali (Bangladesh)
        self.alt_langs = tk.StringVar(value="en-US,hi-IN")      # optional alternatives
        self.punct = tk.BooleanVar(value=True)

        self._build_ui()

    def _build_ui(self):
        s = ttk.Style()
        try: s.theme_use("clam")
        except Exception: pass

        header = ttk.Frame(self.root); header.pack(fill="x", padx=12, pady=(12, 6))
        ttk.Label(header, text="🎤 Simple Google STT (Streaming)", font=("Segoe UI", 16, "bold")).pack(side="left")
        ttk.Label(header, text="Start → Speak → Save transcript", foreground="#6b7280").pack(side="right")

        card = ttk.Frame(self.root); card.pack(fill="x", padx=12, pady=6)

        row1 = ttk.Frame(card); row1.pack(fill="x", padx=4, pady=6)
        ttk.Button(row1, text="Start", command=self.start).pack(side="left", padx=4)
        ttk.Button(row1, text="Stop", command=self.stop).pack(side="left", padx=4)
        ttk.Button(row1, text="Clear", command=self.clear).pack(side="left", padx=4)
        ttk.Button(row1, text="Save .txt", command=self.save_txt).pack(side="left", padx=8)
        ttk.Button(row1, text="Save .docx", command=self.save_docx).pack(side="left", padx=4)

        row2 = ttk.Frame(card); row2.pack(fill="x", padx=4, pady=6)
        ttk.Label(row2, text="Primary language code").pack(side="left", padx=(0,6))
        ttk.Entry(row2, textvariable=self.primary_lang, width=12).pack(side="left", padx=(0,14))
        ttk.Label(row2, text="Alternative language codes (comma-separated)").pack(side="left", padx=(0,6))
        ttk.Entry(row2, textvariable=self.alt_langs, width=28).pack(side="left", padx=(0,14))
        ttk.Checkbutton(row2, text="Automatic punctuation", variable=self.punct).pack(side="left")

        transcript = ttk.Frame(self.root); transcript.pack(fill="both", expand=True, padx=12, pady=(6, 12))
        ttk.Label(transcript, text="Transcript").pack(anchor="w")
        self.text = tk.Text(transcript, wrap="word", font=("Segoe UI", 11), bg="#0b1220", fg="#e5e7eb", insertbackground="#e5e7eb")
        self.text.pack(fill="both", expand=True, padx=4, pady=6)

        footer = ttk.Frame(self.root); footer.pack(fill="x", padx=12, pady=(0, 10))
        self.status = ttk.Label(footer, text="Ready.")
        self.status.pack(side="left")

    def set_status(self, msg):
        # Safe to call from worker threads
        self.root.after(0, lambda: self.status.config(text=msg))

    def append_line(self, msg):
        # Safe to call from worker threads
        def _do():
            self.text.insert(tk.END, msg + "\n")
            self.text.see(tk.END)
        self.root.after(0, _do)

    def clear(self):
        self.text.delete("1.0", tk.END)
        self.lines.clear()

    # ------------- Start / Stop -------------
    def start(self):
        if self.running:
            return
        if gspeech is None:
            self.append_line("[ERROR] google-cloud-speech not installed.")
            return

        self.running = True
        self.stop_event.clear()
        self.lines.clear()
        self.set_status("Starting…")

        try:
            self.mic_device_index, self.input_sr = pick_input_device_and_sr()
        except Exception as e:
            self.append_line(f"[{ts()}] [ERROR] Mic detect: {e}")
            self.stop(); return

        # Build Google configs
        try:
            self.client = gspeech.SpeechClient()
        except Exception as e:
            self.append_line(f"[{ts()}] [ERROR] Google client: {e}")
            self.stop(); return

        primary = (self.primary_lang.get() or "bn-BD").strip()
        alt_codes = [c.strip() for c in (self.alt_langs.get() or "").split(",") if c.strip()]

        self.rec_config = gspeech.RecognitionConfig(
            encoding=gspeech.RecognitionConfig.AudioEncoding.LINEAR16,
            sample_rate_hertz=TARGET_SR,
            language_code=primary,
            alternative_language_codes=alt_codes,
            enable_automatic_punctuation=bool(self.punct.get()),
            model="default",
        )
        self.stream_config = gspeech.StreamingRecognitionConfig(
            config=self.rec_config,
            interim_results=True,
            single_utterance=False,
        )

        # Threads: mic capture & recognizer
        self.cap_thread = threading.Thread(target=self._capture_mic_loop, daemon=True)
        self.rec_thread = threading.Thread(target=self._recognize_loop, daemon=True)
        self.cap_thread.start()
        self.rec_thread.start()

        self.append_line(f"[{ts()}] Started. Mic device #{self.mic_device_index} @ {self.input_sr} Hz")
        self.set_status("Listening…")

    def stop(self):
        if not self.running:
            return
        self.running = False
        self.stop_event.set()
        time.sleep(0.2)
        self.set_status("Stopped.")
        self.append_line(f"[{ts()}] Stopped.")

    # ------------- Mic capture -------------
    def _capture_mic_loop(self):
        """Read mic frames, push to queue as int16 mono."""
        frames = 1024
        stream = None
        try:
            stream = sd.InputStream(samplerate=self.input_sr, channels=1, dtype="int16",
                                    blocksize=frames, device=self.mic_device_index)
            stream.start()
            while not self.stop_event.is_set():
                data, _ = stream.read(frames)
                if data is not None and data.size > 0:
                    # Keep shape (N,) int16
                    self.audio_q.put_nowait(data.reshape(-1).copy())
        except queue.Full:
            pass
        except Exception as e:
            self.append_line(f"[{ts()}] [ERROR] Mic: {e}")
        finally:
            try:
                if stream: stream.stop(); stream.close()
            except Exception:
                pass

    # ------------- Google streaming -------------
    def _audio_request_generator(self):
        """Yields StreamingRecognizeRequest with audio bytes resampled to 16k if needed."""
        while not self.stop_event.is_set():
            try:
                chunk = self.audio_q.get(timeout=0.2)  # int16 mono
            except queue.Empty:
                continue

            # Convert to float32, resample if needed, back to int16 bytes
            f32 = chunk.astype(np.float32) / 32768.0
            if self.input_sr != TARGET_SR:
                f32 = linear_resample_1d(f32, self.input_sr, TARGET_SR)
            i16 = np.clip(f32, -1.0, 1.0)
            i16 = (i16 * 32767.0).astype(np.int16)
            yield gspeech.StreamingRecognizeRequest(audio_content=i16.tobytes())

        # Give Google a moment to drain remaining items
        while not self.audio_q.empty():
            chunk = self.audio_q.get_nowait()
            f32 = chunk.astype(np.float32) / 32768.0
            if self.input_sr != TARGET_SR:
                f32 = linear_resample_1d(f32, self.input_sr, TARGET_SR)
            i16 = np.clip(f32, -1.0, 1.0)
            i16 = (i16 * 32767.0).astype(np.int16)
            yield gspeech.StreamingRecognizeRequest(audio_content=i16.tobytes())

    def _recognize_loop(self):
        try:
            requests = self._audio_request_generator()
            responses = self.client.streaming_recognize(self.stream_config, requests)

            for response in responses:
                if self.stop_event.is_set():
                    break
                for result in response.results:
                    # Show interim in status; only append final to transcript
                    if result.is_final:
                        text = result.alternatives[0].transcript.strip()
                        if text:
                            self.lines.append({"time": ts(), "text": text})
                            self.append_line(f"[{ts()}] {text}")
                            self.set_status("Listening…")
                    else:
                        interim = result.alternatives[0].transcript.strip()
                        if interim:
                            self.set_status(f"Listening… {interim}")
        except Exception as e:
            self.append_line(f"[{ts()}] [ERROR] STT: {e}")
        finally:
            self.set_status("Idle.")

    # ------------- Save actions -------------
    def save_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not path: return
        with open(path, "w", encoding="utf-8") as f:
            for line in self.lines:
                f.write(f"[{line['time']}] {line['text']}\n")
        self.set_status(f"Saved {os.path.basename(path)}")

    def save_docx(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not path: return
        doc = Document(); doc.add_heading("Transcript", level=1)
        for line in self.lines:
            doc.add_paragraph(f"[{line['time']}] {line['text']}")
        doc.save(path); self.set_status(f"Saved {os.path.basename(path)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = GoogleSTTApp(root)
    root.mainloop()
