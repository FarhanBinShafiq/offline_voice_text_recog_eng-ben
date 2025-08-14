import tkinter as tk
from tkinter import ttk, filedialog
import numpy as np
import sounddevice as sd
import threading, queue, time, os, json, datetime

from docx import Document
from faster_whisper import WhisperModel

# Optional system-audio capture (Meeting Mode)
try:
    import soundcard as sc
    HAS_SOUNDCARD = True
except Exception:
    HAS_SOUNDCARD = False


# ---------- Small utils ----------
def ts(): return time.strftime("%H:%M:%S")
def now_tag(): return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
def ensure_dir(path): os.makedirs(path, exist_ok=True); return path


def get_default_input_device_index() -> int:
    """Default input device index, fallback to first input-capable device."""
    try:
        d = sd.default.device
        if isinstance(d, (list, tuple)) and d[0] is not None:
            return int(d[0])
        if isinstance(d, int):
            return d
    except Exception:
        pass
    for i, dev in enumerate(sd.query_devices()):
        if dev.get("max_input_channels", 0) > 0:
            return i
    raise RuntimeError("No input microphone found")


def best_samplerate_for_device(device_index: int) -> int:
    """Prefer 16k for ASR; fallback to 8k (common for BT hands-free); else host default."""
    for sr in (16000, 8000):
        try:
            sd.check_input_settings(device=device_index, samplerate=sr, channels=1)
            return sr
        except Exception:
            continue
    info = sd.query_devices(device_index)
    host = sd.query_hostapis()[info["hostapi"]]
    default_sr = int(host.get("default_samplerate", 16000) or 16000)
    sd.check_input_settings(device=device_index, samplerate=default_sr, channels=1)
    return default_sr


def linear_resample_1d(x: np.ndarray, sr_in: int, sr_out: int) -> np.ndarray:
    """Lightweight linear resampler (mono float32)."""
    if sr_in == sr_out or x.size == 0:
        return x.astype(np.float32, copy=False)
    ratio = float(sr_out) / float(sr_in)
    new_len = max(1, int(round(x.size * ratio)))
    xi = np.linspace(0.0, 1.0, x.size, dtype=np.float32)
    xo = np.linspace(0.0, 1.0, new_len, dtype=np.float32)
    return np.interp(xo, xi, x.astype(np.float32)).astype(np.float32)


# ---------- Ring buffer ----------
class RingBuffer:
    def __init__(self, max_chunks=128):
        self.q = queue.Queue(maxsize=max_chunks)

    def push(self, arr: np.ndarray):
        try:
            self.q.put_nowait(arr)
        except queue.Full:
            try:
                _ = self.q.get_nowait()
                self.q.put_nowait(arr)
            except Exception:
                pass

    def pop_all(self):
        items = []
        try:
            while True:
                items.append(self.q.get_nowait())
        except queue.Empty:
            pass
        if not items:
            return None
        return np.concatenate(items, axis=0)


# ---------- ASR worker ----------
class ASRWorker(threading.Thread):
    """Consumes int16 audio at any SR, resamples to 16k, transcribes with faster-whisper."""
    def __init__(self, rb, on_text, on_info, stop_event,
                 model_name="medium", compute_type="int8",
                 language=None, task="transcribe",
                 input_sr=16000, chunk_secs=6, overlap_secs=0.6):
        super().__init__(daemon=True)
        self.rb = rb
        self.on_text = on_text
        self.on_info = on_info
        self.stop_event = stop_event

        self.model_name = model_name
        self.compute_type = compute_type
        self.language = language
        self.task = task

        self.input_sr = input_sr
        self.target_sr = 16000
        self.stride_in = int(chunk_secs * input_sr)
        self.overlap_in = int(overlap_secs * input_sr)

        self.tail = np.zeros((0, 1), dtype=np.int16)
        self.model = None

    def run(self):
        try:
            self.model = WhisperModel(self.model_name, device="auto", compute_type=self.compute_type)
            self.on_info(f"[{ts()}] ASR loaded: {self.model_name} ({self.compute_type})")
            self.on_info(f"[{ts()}] Input SR: {self.input_sr} Hz → ASR SR: {self.target_sr} Hz")
        except Exception as e:
            self.on_info(f"[{ts()}] [ERROR] Model load failed: {e}")
            return

        while not self.stop_event.is_set():
            time.sleep(0.02)
            buf = self.rb.pop_all()
            if buf is None or buf.size == 0:
                continue

            if self.tail.size:
                buf = np.concatenate([self.tail, buf], axis=0)

            if buf.shape[0] < self.stride_in:
                self.tail = buf
                continue

            chunk = buf[:self.stride_in]
            keep_from = max(0, self.stride_in - self.overlap_in)
            self.tail = buf[keep_from:]

            # int16 -> float32 [-1,1] @ input_sr
            audio_f32 = (chunk.astype(np.float32) / 32768.0).squeeze(axis=-1)
            # resample to 16k for ASR
            audio_16k = linear_resample_1d(audio_f32, self.input_sr, self.target_sr)

            try:
                segments, _ = self.model.transcribe(
                    audio_16k,
                    language=self.language,               # None = auto
                    task=self.task,                       # 'transcribe' or 'translate'
                    vad_filter=False,
                    beam_size=1, best_of=1,
                    no_speech_threshold=0.2,
                    condition_on_previous_text=True
                )
                out = []
                srt_segments = []
                for seg in segments:
                    text = (seg.text or "").strip()
                    if text:
                        out.append(text)
                        srt_segments.append({
                            "start": float(seg.start or 0.0),
                            "end": float(seg.end or 0.0),
                            "text": text
                        })
                if out:
                    self.on_text(" ".join(out), srt_segments)
            except Exception as e:
                self.on_info(f"[{ts()}] [ERROR] ASR failed: {e}")


# ---------- GUI ----------
class MeetingTranscriberApp:
    DEFAULT_SR = 16000  # target SR for ASR (resampled)

    def __init__(self, root):
        self.root = root
        self.root.title("Meeting Transcriber (Offline)")
        self.root.geometry("1020x700")
        self.root.configure(bg="#0b1220")

        # state
        self.running = False
        self.stop_event = threading.Event()
        self.rb = RingBuffer()

        self.lines = []     # [{"time": "...", "text": "..."}]
        self.segments = []  # [{"start": s, "end": e, "text": "..."}]
        self.last_autosave = time.time()
        self.meeting_dir = ensure_dir(os.path.join(os.getcwd(), f"meeting_{now_tag()}"))

        # UI state
        self.model_name = tk.StringVar(value="medium")
        self.compute_type = tk.StringVar(value="int8")
        self.lang = tk.StringVar(value="Auto (detect)")
        self.translate = tk.BooleanVar(value=False)
        self.chunk_secs = tk.IntVar(value=6)
        self.overlap = tk.DoubleVar(value=0.6)
        self.meeting_mode = tk.BooleanVar(value=False)

        # capture settings (resolved at start)
        self.mic_device_index = None
        self.input_sr = self.DEFAULT_SR

        self._build_ui()

    # ----- UI -----
    def _build_ui(self):
        s = ttk.Style()
        try: s.theme_use("clam")
        except Exception: pass
        s.configure("TFrame", background="#0b1220")
        s.configure("Card.TFrame", background="#0f1626", relief="groove", borderwidth=1)
        s.configure("TLabel", foreground="#e5e7eb", background="#0b1220")
        s.configure("Title.TLabel", font=("Segoe UI", 18, "bold"), foreground="#e5e7eb", background="#0b1220")
        s.configure("Sub.TLabel", font=("Segoe UI", 10), foreground="#9ca3af", background="#0b1220")
        s.configure("TButton", font=("Segoe UI", 10))
        s.configure("Accent.TButton", font=("Segoe UI", 10, "bold"))
        s.map("Accent.TButton", background=[("!disabled", "#22c55e"), ("disabled", "#14532d")], foreground=[("!disabled", "white")])
        s.configure("Danger.TButton", font=("Segoe UI", 10, "bold"))
        s.map("Danger.TButton", background=[("!disabled", "#ef4444"), ("disabled", "#7f1d1d")], foreground=[("!disabled", "white")])
        s.configure("TProgressbar", troughcolor="#111827", background="#3b82f6")

        header = ttk.Frame(self.root, style="TFrame")
        header.pack(fill="x", padx=16, pady=(16, 8))
        ttk.Label(header, text="🎤 Meeting Transcriber (Offline)", style="Title.TLabel").pack(side="left")
        ttk.Label(header, text="Faster-Whisper • Autosave • TXT/DOCX/JSON/SRT", style="Sub.TLabel").pack(side="right")

        card = ttk.Frame(self.root, style="Card.TFrame")
        card.pack(fill="x", padx=16, pady=8)

        row1 = ttk.Frame(card, style="TFrame")
        row1.pack(fill="x", padx=12, pady=10)
        self.btn_start = ttk.Button(row1, text="Start Meeting", style="Accent.TButton", command=self.start)
        self.btn_stop  = ttk.Button(row1, text="Stop", style="Danger.TButton", state="disabled", command=self.stop)
        self.btn_clear = ttk.Button(row1, text="Clear", command=self.clear)
        self.btn_save_txt  = ttk.Button(row1, text="Save .txt", command=self.save_txt)
        self.btn_save_docx = ttk.Button(row1, text="Save .docx", command=self.save_docx)
        self.btn_save_json = ttk.Button(row1, text="Save .json", command=self.save_json)
        self.btn_save_srt  = ttk.Button(row1, text="Save .srt", command=self.save_srt)
        for w in (self.btn_start, self.btn_stop, self.btn_clear, self.btn_save_txt, self.btn_save_docx, self.btn_save_json, self.btn_save_srt):
            w.pack(side="left", padx=6)

        row2 = ttk.Frame(card, style="TFrame"); row2.pack(fill="x", padx=12, pady=6)
        ttk.Label(row2, text="Model").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.model_name, values=["base","small","medium","large-v3"], width=12, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Label(row2, text="Compute").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.compute_type, values=["int8","float16","float32"], width=10, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Label(row2, text="Language").pack(side="left", padx=(0, 6))
        ttk.Combobox(row2, textvariable=self.lang, values=["Auto (detect)","English (en)","Hindi (hi)","Japanese (ja)","Spanish (es)"], width=18, state="readonly").pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row2, text="Translate to English", variable=self.translate).pack(side="left", padx=(0, 16))
        ttk.Checkbutton(row2, text=f"Meeting Mode (capture speakers){'' if HAS_SOUNDCARD else ' [install soundcard]'}",
                        variable=self.meeting_mode, state=("!disabled" if HAS_SOUNDCARD else "disabled")).pack(side="left")

        # Row 3: (labels first, then attach callbacks to avoid early-callback crash)
        row3 = ttk.Frame(card, style="TFrame"); row3.pack(fill="x", padx=12, pady=(0, 12))

        ttk.Label(row3, text="Chunk (s)").pack(side="left", padx=(0, 6))
        self.scale_chunk = ttk.Scale(row3, from_=3, to=12, orient="horizontal")
        self.scale_chunk.pack(side="left", padx=(0, 16))
        self.lbl_chunk = ttk.Label(row3, text=f"{self.chunk_secs.get()}s")
        self.lbl_chunk.pack(side="left", padx=(0, 16))
        self.scale_chunk.set(self.chunk_secs.get())
        self.scale_chunk.configure(command=self._on_chunk_change)

        ttk.Label(row3, text="Overlap").pack(side="left", padx=(0, 6))
        self.scale_olap = ttk.Scale(row3, from_=0.2, to=0.9, orient="horizontal")
        self.scale_olap.pack(side="left", padx=(0, 16))
        self.lbl_olap = ttk.Label(row3, text=f"{self.overlap.get():.1f}")
        self.lbl_olap.pack(side="left", padx=(0, 16))
        self.scale_olap.set(self.overlap.get())
        self.scale_olap.configure(command=self._on_overlap_change)

        transcript = ttk.Frame(self.root, style="Card.TFrame")
        transcript.pack(fill="both", expand=True, padx=16, pady=(8, 16))
        ttk.Label(transcript, text="Transcript", style="Sub.TLabel").pack(anchor="w", padx=12, pady=(8, 0))
        self.text = tk.Text(transcript, wrap="word", font=("Segoe UI", 11),
                            bg="#09101d", fg="#e5e7eb", insertbackground="#e5e7eb", relief="flat")
        self.text.pack(fill="both", expand=True, padx=12, pady=12)

        footer = ttk.Frame(self.root, style="TFrame"); footer.pack(fill="x", padx=16, pady=(0, 12))
        self.status = ttk.Label(footer, text="Ready.", style="Sub.TLabel"); self.status.pack(side="left")

    def _on_chunk_change(self, *_):
        v = int(float(self.scale_chunk.get())); self.chunk_secs.set(v)
        if hasattr(self, "lbl_chunk"): self.lbl_chunk.config(text=f"{v}s")

    def _on_overlap_change(self, *_):
        v = float(self.scale_olap.get()); self.overlap.set(round(v, 1))
        if hasattr(self, "lbl_olap"): self.lbl_olap.config(text=f"{self.overlap.get():.1f}")

    def set_status(self, msg): self.status.config(text=msg)

    # ----- Start/Stop -----
    def start(self):
        if self.running: return
        self.running = True
        self.stop_event.clear()
        self.btn_start.config(state="disabled"); self.btn_stop.config(state="normal")
        self.lines.clear(); self.segments.clear()
        self.set_status("Starting…")

        # resolve mic + SR (for mic mode)
        try:
            self.mic_device_index = get_default_input_device_index()
            self.input_sr = best_samplerate_for_device(self.mic_device_index)
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Mic detect: {e}")
            self.stop(); return

        # language / task
        lang = None
        if self.lang.get() != "Auto (detect)":
            lang = self.lang.get().split("(")[-1].split(")")[0]
        task = "translate" if self.translate.get() else "transcribe"

        # ASR worker
        self.asr_thread = ASRWorker(
            rb=self.rb, on_text=self._on_text, on_info=self._append_line, stop_event=self.stop_event,
            model_name=self.model_name.get(), compute_type=self.compute_type.get(),
            language=lang, task=task,
            input_sr=(self.input_sr if not self.meeting_mode.get() else 48000),  # meeting loopback often 48k
            chunk_secs=self.chunk_secs.get(), overlap_secs=self.overlap.get()
        )
        self.asr_thread.start()

        # recorder thread
        rec_target = self._record_loop_meeting if (self.meeting_mode.get() and HAS_SOUNDCARD) else self._record_loop_mic
        self.rec_thread = threading.Thread(target=rec_target, daemon=True)
        self.rec_thread.start()
        self._append_line(f"[{ts()}] Meeting started. Saving in: {self.meeting_dir}")

    def stop(self):
        if not self.running: return
        self.running = False
        self.stop_event.set()
        self.btn_start.config(state="normal"); self.btn_stop.config(state="disabled")
        self.set_status("Stopping…"); time.sleep(0.3); self.set_status("Stopped.")

    def clear(self):
        self.text.delete("1.0", tk.END)
        self.lines.clear(); self.segments.clear()

    # ----- Recording -----
    def _record_loop_mic(self):
        """Capture laptop/default mic at the device's supported SR (8k/16k/etc)."""
        frames = 1024
        try:
            stream = sd.InputStream(samplerate=self.input_sr, channels=1, dtype="int16",
                                    blocksize=frames, device=self.mic_device_index)
            stream.start()
            self._append_line(f"[{ts()}] Mic capture: device #{self.mic_device_index} @ {self.input_sr} Hz")
            self.set_status("Listening (mic)…")

            accum = []
            target = int(self.chunk_secs.get() * self.input_sr)
            while not self.stop_event.is_set():
                data, _ = stream.read(frames)
                if data is not None and data.size > 0:
                    accum.append(data)
                total = sum(x.shape[0] for x in accum)
                if total >= target // 3:
                    self.rb.push(np.concatenate(accum, axis=0)); accum = []
                self._autosave_if_needed()
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Mic: {e}")
        finally:
            try: stream.stop(); stream.close()
            except Exception: pass

    def _record_loop_meeting(self):
        """Capture system audio (speakers) + mic mix using soundcard at 48 kHz, then queue int16."""
        if not HAS_SOUNDCARD:
            self._append_line(f"[{ts()}] [WARN] soundcard not installed; fallback to mic.")
            self._record_loop_mic(); return

        sr = 48000  # robust for loopback; ASR worker will resample to 16k
        block = 1024
        try:
            speaker = sc.default_speaker(); mic = sc.default_microphone()
            self._append_line(f"[{ts()}] Meeting mode: loopback @ {sr} Hz")
            self.set_status("Listening (meeting mode)…")
            with mic.recorder(samplerate=sr, channels=1, blocksize=block) as mic_rec, \
                 speaker.recorder(samplerate=sr, channels=1, blocksize=block) as spk_rec:
                accum = []
                target = int(self.chunk_secs.get() * sr)
                while not self.stop_event.is_set():
                    mic_data = mic_rec.record(numframes=block)      # float32 [-1,1]
                    spk_data = spk_rec.record(numframes=block)      # float32 [-1,1]
                    mixed = np.clip(mic_data + spk_data, -1.0, 1.0) # mix
                    i16 = (mixed * 32767.0).astype(np.int16)        # → int16
                    accum.append(i16)
                    total = sum(x.shape[0] for x in accum)
                    if total >= target // 3:
                        self.rb.push(np.concatenate(accum, axis=0)); accum = []
                    self._autosave_if_needed()
        except Exception as e:
            self._append_line(f"[{ts()}] [ERROR] Meeting mode: {e}")
            self._record_loop_mic()

    # ----- Callbacks from ASR -----
    def _on_text(self, text_line, srt_segments):
        self.text.insert(tk.END, f"[{ts()}] {text_line}\n")
        self.text.see(tk.END)
        self.lines.append({"time": ts(), "text": text_line})
        self.segments.extend(srt_segments)
        self._autosave_if_needed()

    def _append_line(self, text):
        self.text.insert(tk.END, text + "\n")
        self.text.see(tk.END)

    # ----- Autosave -----
    def _autosave_if_needed(self):
        if time.time() - self.last_autosave > 60:
            self._save_txt(os.path.join(self.meeting_dir, "autosave.txt"))
            self._save_json(os.path.join(self.meeting_dir, "autosave.json"))
            self._save_srt(os.path.join(self.meeting_dir, "autosave.srt"))
            self.last_autosave = time.time()

    # ----- Save actions -----
    def save_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not path: return
        self._save_txt(path)

    def save_docx(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not path: return
        doc = Document(); doc.add_heading("Meeting Transcript", level=1)
        for line in self.lines:
            doc.add_paragraph(f"[{line['time']}] {line['text']}")
        doc.save(path); self.set_status(f"Saved {os.path.basename(path)}")

    def save_json(self):
        path = filedialog.asksaveasfilename(defaultextension=".json")
        if not path: return
        self._save_json(path)

    def save_srt(self):
        path = filedialog.asksaveasfilename(defaultextension=".srt")
        if not path: return
        self._save_srt(path)

    # ----- Save helpers -----
    def _save_txt(self, path):
        with open(path, "w", encoding="utf-8") as f:
            for line in self.lines:
                f.write(f"[{line['time']}] {line['text']}\n")
        self.set_status(f"Saved {os.path.basename(path)}")

    def _save_json(self, path):
        payload = {
            "created": datetime.datetime.now().isoformat(),
            "model": self.model_name.get(),
            "compute": self.compute_type.get(),
            "language": self.lang.get(),
            "translate_to_english": self.translate.get(),
            "lines": self.lines,
            "segments": self.segments,
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        self.set_status(f"Saved {os.path.basename(path)}")

    def _save_srt(self, path):
        def fmt_time(sec):
            h = int(sec // 3600); m = int((sec % 3600) // 60); s = int(sec % 60)
            ms = int((sec - int(sec)) * 1000)
            return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

        out, idx = [], 1
        for seg in self.segments:
            start = max(0.0, float(seg.get("start", 0.0)))
            end   = max(start, float(seg.get("end", start + 0.1)))
            text  = (seg.get("text") or "").strip()
            if not text: continue
            out.append(f"{idx}\n{fmt_time(start)} --> {fmt_time(end)}\n{text}\n")
            idx += 1
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(out))
        self.set_status(f"Saved {os.path.basename(path)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = MeetingTranscriberApp(root)
    root.mainloop()
